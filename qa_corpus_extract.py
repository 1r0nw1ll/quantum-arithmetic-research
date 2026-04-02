#!/usr/bin/env python3
"""
qa_corpus_extract.py — QA Corpus Text Extraction (Step 1: no-OCR pass)

Extracts machine-readable text from all DOCX and good-text PDF files in the
QA/Pythagoras corpus. Outputs one Markdown file per book into qa_corpus_text/.

Only processes files rated EASY in the audit (docx_extractable or good/partial PDF).
Image-only PDFs are logged as TODO for OCR pass.

Usage:
  python qa_corpus_extract.py
  python qa_corpus_extract.py --out-dir /path/to/output
  python qa_corpus_extract.py --dry-run
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Optional

CORPUS_ROOT = Path("/home/player2/Desktop/files/quantum_pythagoras-text/quantum_pythagoras-text")
OUT_DIR     = Path("/home/player2/signal_experiments/qa_corpus_text")
AUDIT_JSON  = Path("/home/player2/signal_experiments/qa_corpus_audit_report.json")

# ── DOCX extraction ───────────────────────────────────────────────────────────

def extract_docx(path: Path) -> Optional[str]:
    """Extract plain text from DOCX. Returns None on failure."""
    try:
        import docx as python_docx
        doc = python_docx.Document(str(path))
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Detect headings by style name
                style = para.style.name if para.style else ""
                if "Heading" in style:
                    level = re.search(r"\d+", style)
                    level = int(level.group()) if level else 1
                    lines.append(f"\n{'#' * level} {text}\n")
                else:
                    lines.append(text)

        # Tables
        for table in doc.tables:
            lines.append("\n[TABLE]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
            lines.append("[/TABLE]\n")

        return "\n".join(lines)
    except Exception as e:
        return f"[EXTRACTION ERROR: {e}]"


# ── PDF text extraction ───────────────────────────────────────────────────────

def extract_pdf(path: Path, quality: str) -> Optional[str]:
    """Extract text from PDF using pymupdf. Returns None if image-only."""
    if quality == "none":
        return None
    try:
        import fitz
        doc = fitz.open(str(path))
        pages = []
        for i, page in enumerate(doc):
            text = page.get_text("text").strip()
            if text:
                pages.append(f"\n<!-- page {i+1} -->\n{text}")
        doc.close()
        if not pages:
            return None
        return "\n".join(pages)
    except Exception as e:
        return f"[EXTRACTION ERROR: {e}]"


# ── RTF extraction ────────────────────────────────────────────────────────────

def extract_rtf(path: Path) -> str:
    """Basic RTF text extraction by stripping control words."""
    try:
        raw = path.read_bytes().decode("latin-1", errors="replace")
        # Strip RTF control sequences
        text = re.sub(r"\\[a-z]+[-\d]*\s?", " ", raw)
        text = re.sub(r"[{}\\]", "", text)
        text = re.sub(r"\s+", " ", text).strip()
        return text
    except Exception as e:
        return f"[EXTRACTION ERROR: {e}]"


# ── slug generation ───────────────────────────────────────────────────────────

def make_slug(series: str, file_type: str, filename: str) -> str:
    """Generate a clean output filename."""
    series_clean = re.sub(r"[^a-zA-Z0-9_-]", "_", series).lower()
    type_tag = file_type.lower()
    # Strip common suffixes from filename
    name = Path(filename).stem
    name = re.sub(r"[-_\s]+", "_", name).lower()
    name = name[:40]
    return f"{series_clean}__{name}__{type_tag}.md"


# ── main ──────────────────────────────────────────────────────────────────────

def run_extraction(out_dir: Path, dry_run: bool = False) -> dict:
    if not AUDIT_JSON.exists():
        print("ERROR: audit report not found. Run qa_corpus_audit.py first.")
        sys.exit(1)

    audit = json.loads(AUDIT_JSON.read_text())
    files = audit["files"]

    if not dry_run:
        out_dir.mkdir(parents=True, exist_ok=True)

    stats = {"extracted": 0, "skipped_ocr": 0, "skipped_other": 0, "errors": 0}
    ocr_needed: list[str] = []
    extracted_files: list[str] = []

    print(f"Output dir: {out_dir}")
    print(f"Processing {len(files)} files...\n")

    for f in files:
        rel  = f["rel_path"]
        diff = f["ocr_difficulty"]
        ftype = f["file_type"]
        series = f["series"]
        quality = f.get("text_layer_quality", "none")
        docx_ok = f.get("docx_extractable", False)
        docx_chars = f.get("docx_chars", 0)

        abs_path = CORPUS_ROOT / rel

        # Skip covers and tiny files
        if diff == "skip":
            stats["skipped_other"] += 1
            continue

        # Skip macOS resource forks
        if Path(rel).name.startswith("._"):
            continue

        # Skip duplicate copies in same series (prefer main DOCX over "(1)" and "(2)")
        if re.search(r"\(\d+\)\.docx$", rel):
            print(f"  → skip duplicate: {Path(rel).name}")
            stats["skipped_other"] += 1
            continue

        # Determine extraction method
        text: Optional[str] = None
        method = ""

        if ftype == "docx" and docx_ok:
            text = extract_docx(abs_path)
            method = "docx"
        elif ftype == "docx" and not docx_ok:
            # Try PDF fallback
            print(f"  ~ DOCX empty, using PDF fallback: {Path(rel).name}")
            # Look for paired PDF
            paired_pdf = abs_path.with_suffix(".pdf")
            if paired_pdf.exists():
                pdf_quality = "partial"  # conservative
                text = extract_pdf(paired_pdf, pdf_quality)
                method = "pdf_fallback"
                rel = str(paired_pdf.relative_to(CORPUS_ROOT))
        elif ftype == "pdf":
            if quality in ("good", "partial"):
                text = extract_pdf(abs_path, quality)
                method = "pdf"
            else:
                # Image-only
                ocr_needed.append(rel)
                stats["skipped_ocr"] += 1
                print(f"  ✗ OCR needed: {Path(rel).name}")
                continue
        elif ftype == "rtf":
            text = extract_rtf(abs_path)
            method = "rtf"

        if text is None or (isinstance(text, str) and len(text) < 50):
            print(f"  ~ no text extracted: {Path(rel).name} [{method}]")
            if quality == "none" or not docx_ok:
                ocr_needed.append(rel)
                stats["skipped_ocr"] += 1
            else:
                stats["skipped_other"] += 1
            continue

        if "[EXTRACTION ERROR" in text:
            print(f"  ✗ error: {Path(rel).name}: {text[:80]}")
            stats["errors"] += 1
            continue

        # Build output markdown
        slug = make_slug(series, method, Path(rel).name)
        char_count = len(text)

        header = f"""---
source: {rel}
series: {series}
method: {method}
chars: {char_count}
extracted: 2026-03-26
---

# {series} — {Path(rel).stem}

"""
        full_text = header + text

        out_path = out_dir / slug
        if not dry_run:
            out_path.write_text(full_text, encoding="utf-8")

        print(f"  ✓ {series:12s} [{method:12s}] {char_count:>8,} chars → {slug}")
        stats["extracted"] += 1
        extracted_files.append(str(out_path))

    # Summary
    print(f"\n{'='*60}")
    print(f"EXTRACTION SUMMARY")
    print(f"{'='*60}")
    print(f"  Extracted:      {stats['extracted']} files")
    print(f"  OCR needed:     {stats['skipped_ocr']} files")
    print(f"  Skipped other:  {stats['skipped_other']} files")
    print(f"  Errors:         {stats['errors']} files")

    if ocr_needed:
        print(f"\nOCR TODO ({len(ocr_needed)} files):")
        for f in ocr_needed:
            print(f"  {f}")

    # Total chars
    if not dry_run:
        total_chars = 0
        for fp in extracted_files:
            p = Path(fp)
            if p.exists():
                total_chars += len(p.read_text(encoding="utf-8"))
        print(f"\nTotal corpus text extracted: {total_chars:,} chars")

    return {"stats": stats, "ocr_needed": ocr_needed, "extracted": extracted_files}


# ── CLI ───────────────────────────────────────────────────────────────────────

def main(argv=None):
    parser = argparse.ArgumentParser(description="QA Corpus Text Extraction (no-OCR pass)")
    parser.add_argument("--out-dir", type=Path, default=OUT_DIR)
    parser.add_argument("--dry-run", action="store_true",
                        help="Show what would be extracted without writing files")
    args = parser.parse_args(argv)

    run_extraction(args.out_dir, args.dry_run)


if __name__ == "__main__":
    main()
